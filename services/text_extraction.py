import asyncio
import tempfile
import os
import subprocess
import glob
from typing import Optional
from PIL import Image
import pytesseract

# Auto-configure tesseract path
def find_tesseract_path():
    """Automatically find tesseract executable path."""
    # Common paths where tesseract might be installed
    common_paths = [
        '/usr/bin/tesseract',
        '/usr/local/bin/tesseract',
        '/opt/homebrew/bin/tesseract',
        '/nix/store/*/bin/tesseract'  # Nix path pattern
    ]
    
    # Try to find using 'which' command first
    try:
        result = subprocess.run(['which', 'tesseract'], capture_output=True, text=True)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except:
        pass
    
    # Try common paths
    for path in common_paths:
        if '*' in path:
            # Handle Nix store paths with glob
            import glob
            matches = glob.glob(path)
            if matches:
                return matches[0]
        elif os.path.exists(path):
            return path
    
    return None

# Configure tesseract path
tesseract_path = find_tesseract_path()
if tesseract_path:
    pytesseract.pytesseract.tesseract_cmd = tesseract_path
    print(f"Tesseract found at: {tesseract_path}")
else:
    print("Tesseract not found in common paths")

# Set partition to None to force fallback
partition = None

import pdfplumber
from PyPDF2 import PdfReader
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class TextExtractionService:
    """Service for extracting text from documents using unstructured library."""

    def __init__(self):
        """Initialize the text extraction service."""
        pass

    async def extract_with_unstructured(self, file_content: bytes, filename: str) -> str:
        """
        Extract text using unstructured library with fallback.

        Args:
            file_content: Raw file content
            filename: Original filename

        Returns:
            Extracted text with layout preservation
        """
        try:
            # If unstructured is not available or fails, use fallback
            if partition is None:
                return await self._fallback_extraction(file_content, filename)
                
            # Write content to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=self._get_file_extension(filename)) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                # Run extraction in thread pool to avoid blocking
                loop = asyncio.get_event_loop()
                elements = await loop.run_in_executor(None, partition, temp_file_path)

                # Clean up temporary file
                os.unlink(temp_file_path)

                # Convert elements to text while preserving structure
                extracted_text = ""
                for element in elements:
                    text = str(element)
                    # Add spacing based on element type to preserve layout
                    if hasattr(element, 'category'):
                        if element.category in ['Title', 'Header']:
                            extracted_text += f"\n# {text}\n\n"
                        elif element.category == 'Table':
                            extracted_text += f"\n{text}\n\n"
                        elif element.category == 'ListItem':
                            extracted_text += f"- {text}\n"
                        else:
                            extracted_text += f"{text}\n\n"
                    else:
                        extracted_text += f"{text}\n\n"

                return extracted_text.strip()
                
            except Exception:
                # Clean up temporary file if error occurs
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                # Fallback to alternative extraction
                return await self._fallback_extraction(file_content, filename)

        except Exception as e:
            # Final fallback
            return await self._fallback_extraction(file_content, filename)

    async def extract_with_marker(self, file_content: bytes, filename: str) -> str:
        """
        Fallback to unstructured for marker strategy.

        Args:
            file_content: Raw file content
            filename: Original filename

        Returns:
            Extracted text using unstructured
        """
        # Since marker has issues, fallback to unstructured
        return await self.extract_with_unstructured(file_content, filename)

    async def _fallback_extraction(self, file_content: bytes, filename: str) -> str:
        """
        Fallback text extraction for when unstructured fails.
        
        Args:
            file_content: Raw file content
            filename: Original filename
            
        Returns:
            Extracted text using alternative methods
        """
        try:
            file_ext = self._get_file_extension(filename).lower()
            
            if file_ext == '.pdf':
                return await self._extract_pdf_fallback(file_content)
            elif file_ext in ['.docx', '.doc']:
                return await self._extract_word_document(file_content)
            elif file_ext in ['.txt', '.md']:
                return file_content.decode('utf-8', errors='ignore')
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif', '.webp']:
                return await self._extract_image_text(file_content)
            else:
                # For other file types, try to decode as text
                return file_content.decode('utf-8', errors='ignore')
                
        except Exception:
            return "Error: Could not extract text from the document."
    
    async def _extract_pdf_fallback(self, file_content: bytes) -> str:
        """Extract text from PDF using pdfplumber and PyPDF2 as fallback."""
        try:
            # Try pdfplumber first
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                loop = asyncio.get_event_loop()
                text = await loop.run_in_executor(None, self._pdfplumber_extract, temp_file_path)
                os.unlink(temp_file_path)
                return text
            except Exception:
                # Fallback to PyPDF2
                try:
                    text = await loop.run_in_executor(None, self._pypdf2_extract, temp_file_path)
                    os.unlink(temp_file_path)
                    return text
                except Exception:
                    os.unlink(temp_file_path)
                    return "Error: Could not extract text from PDF."
                    
        except Exception:
            return "Error: Could not process PDF file."
    
    def _pdfplumber_extract(self, file_path: str) -> str:
        """Extract text using pdfplumber."""
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            return text.strip()
    
    def _pypdf2_extract(self, file_path: str) -> str:
        """Extract text using PyPDF2."""
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n\n"
            return text.strip()

    async def _extract_word_document(self, file_content: bytes) -> str:
        """Extract text from Word document using python-docx."""
        try:
            if not DOCX_AVAILABLE:
                return "Error: python-docx library not available for Word document processing."
            
            # Write content to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                # Extract text using python-docx
                loop = asyncio.get_event_loop()
                extracted_text = await loop.run_in_executor(None, self._extract_docx_text, temp_file_path)
                
                # Clean up
                os.unlink(temp_file_path)
                
                if extracted_text.strip():
                    return extracted_text.strip()
                else:
                    return "No text found in the Word document."
                    
            except Exception as e:
                # Clean up temporary file if error occurs
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                return f"Error extracting text from Word document: {str(e)}"
                
        except Exception as e:
            return f"Error processing Word document: {str(e)}"

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from .docx file using python-docx."""
        doc = Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables with better formatting
        for table_idx, table in enumerate(doc.tables):
            table_text = []
            table_text.append(f"\n--- Table {table_idx + 1} ---")
            
            for row_idx, row in enumerate(table.rows):
                row_cells = []
                for cell in row.cells:
                    # Clean cell text and handle multi-line content
                    cell_text = cell.text.strip().replace('\n', ' ').replace('\r', '')
                    if cell_text:
                        row_cells.append(cell_text)
                    else:
                        row_cells.append("")  # Keep empty cells for structure
                
                if any(cell.strip() for cell in row_cells):  # Only add non-empty rows
                    # Format as table row
                    if row_idx == 0:  # Header row
                        table_text.append(" | ".join(row_cells))
                        table_text.append("-" * len(" | ".join(row_cells)))  # Add separator
                    else:
                        table_text.append(" | ".join(row_cells))
            
            if len(table_text) > 1:  # Only add if table has content
                text_parts.extend(table_text)
                text_parts.append("")  # Add spacing after table
        
        return "\n".join(text_parts)

    async def _extract_image_text(self, file_content: bytes) -> str:
        """Extract text from image using OCR."""
        try:
            # Check if tesseract is available
            try:
                pytesseract.get_tesseract_version()
            except Exception as e:
                return f"Error: Tesseract OCR is not available. Current path: {pytesseract.pytesseract.tesseract_cmd}. Error: {str(e)}"
            
            # Write content to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                # Open image with PIL
                image = Image.open(temp_file_path)
                
                # Convert to RGB if necessary
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                # Run OCR in thread pool with basic English first
                loop = asyncio.get_event_loop()
                try:
                    # Try with Vietnamese + English
                    extracted_text = await loop.run_in_executor(
                        None, 
                        pytesseract.image_to_string, 
                        image,
                        'vie+eng'
                    )
                except Exception:
                    # Fallback to English only
                    extracted_text = await loop.run_in_executor(
                        None, 
                        pytesseract.image_to_string, 
                        image,
                        'eng'
                    )
                
                # Clean up
                os.unlink(temp_file_path)
                
                if extracted_text.strip():
                    return extracted_text.strip()
                else:
                    return "No text found in the image."
                    
            except Exception as e:
                # Clean up temporary file if error occurs
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                return f"Error extracting text from image: {str(e)}"
                
        except Exception as e:
            return f"Error processing image file: {str(e)}"

    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename."""
        return os.path.splitext(filename)[1] or '.txt'